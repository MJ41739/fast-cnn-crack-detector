import docx
import os

def update_docx_report():
    doc_path = r"c:\Users\Mayur Jadhav\OneDrive\Desktop\FASTCNN\outputs\model_specifications_and_evaluation_report.docx"
    if not os.path.exists(doc_path):
        print(f"Error: Word document not found at {doc_path}")
        return
        
    print(f"Loading document: {doc_path}")
    doc = docx.Document(doc_path)
    
    # 1. Update Table 2 (comparison table) which is doc.tables[1]
    table2 = doc.tables[1]
    already_exists = False
    for row in table2.rows:
        if "rcnn" in row.cells[0].text.lower():
            already_exists = True
            break
            
    if not already_exists:
        row_cells = table2.add_row().cells
        row_cells[0].text = 'Recurrent CNN (RCNN)'
        row_cells[1].text = '93.32%'
        row_cells[2].text = '88.60%'
        row_cells[3].text = '99.43%'
        row_cells[4].text = '93.70%'
        row_cells[5].text = '0.9976'
        row_cells[6].text = '0.17 ms'
        print("Added RCNN to Table 2 (Quantitative metrics).")
    else:
        print("RCNN already exists in Table 2.")
        
    # 2. Insert section 2.4 Model D: Recurrent CNN (RCNN)
    target_p = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("3. Quantitative Model"):
            target_p = p
            break
            
    if target_p:
        text_to_insert = [
            "2.4 Model D: Recurrent CNN (RCNN)",
            "The Recurrent Convolutional Neural Network (RCNN) is a specialized neural network architecture designed to enhance feature representations by incorporating recurrent connections within the convolutional layers. Based on the architecture by Liang & Hu (2015), the RCNN utilizes Recurrent Convolutional Layers (RCL) where the state of each unit evolves over discrete time steps. This allows the model to integrate contextual information and expand its receptive field dynamically without increasing the parameter footprint.",
            "For this binary classification task, the Recurrent CNN is configured with a base convolution layer that projects the 3-channel input to 32 feature maps with stride 2. This is followed by four RCL blocks with time step T = 3, configured as follows: Block 1 (64 channels), Block 2 (128 channels), Block 3 (256 channels), and Block 4 (512 channels). Max pooling (2x2) and dropout are applied between the blocks. A Global Average Pooling layer reduces the spatial features to 512 dimensions, which are passed to a dense head of 128 units, followed by a final linear layer outputting the classification logit.",
            "Total Trainable Parameters: 4,768,801 parameters. The RCNN is initialized from scratch and optimized using AdamW under mixed precision."
        ]
        section_exists = False
        for p in doc.paragraphs:
            if "2.4 Model D: Recurrent CNN" in p.text:
                section_exists = True
                break
                
        if not section_exists:
            for text in text_to_insert:
                new_p = target_p.insert_paragraph_before(text)
                if text.startswith("2.4"):
                    new_p.style = doc.styles['Heading 2'] if 'Heading 2' in doc.styles else target_p.style
                else:
                    new_p.style = doc.styles['Normal'] if 'Normal' in doc.styles else target_p.style
            print("Inserted Section 2.4 (RCNN specifications).")
        else:
            print("Section 2.4 already exists.")
            
    # 3. Add Section 4.4 Algorithm 4: Recurrent CNN (RCNN)
    target_p5 = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("5. Explainability"):
            target_p5 = p
            break
            
    if target_p5:
        text_to_insert_figs = [
            "4.4 Algorithm 4: Recurrent CNN (RCNN)",
            "Figure 1. Training & Validation curves for Recurrent CNN (RCNN)",
            "Figure 2. Test Confusion Matrix for Recurrent CNN (RCNN)",
            "Figure 3. ROC-AUC curve for Recurrent CNN (RCNN)"
        ]
        section_figs_exists = False
        for p in doc.paragraphs:
            if "4.4 Algorithm 4: Recurrent CNN" in p.text:
                section_figs_exists = True
                break
                
        if not section_figs_exists:
            for text in text_to_insert_figs:
                new_p = target_p5.insert_paragraph_before(text)
                if text.startswith("4.4"):
                    new_p.style = doc.styles['Heading 2'] if 'Heading 2' in doc.styles else target_p5.style
                else:
                    new_p.style = doc.styles['Normal'] if 'Normal' in doc.styles else target_p5.style
            print("Inserted Section 4.4 (RCNN evaluation figures captions).")
        else:
            print("Section 4.4 already exists.")
            
    doc.save(doc_path)
    print("Updated docx successfully!")

if __name__ == "__main__":
    update_docx_report()
